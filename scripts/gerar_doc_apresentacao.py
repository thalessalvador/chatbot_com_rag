"""Gera documento-base para apresentação oral (projeto final PLN / RAG)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "Apresentacao_Projeto_Final_PLN.docx"


def add_heading_r(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()
    t = doc.add_heading(
        "Apresentação — Projeto Final PLN (IFG)",
        0,
    )
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_p(
        doc,
        "Documento-base para slides e fala (~15 minutos). Público: banca (docentes) e turma de pós-graduação em Computação.",
    )
    add_p(doc, "")

    add_heading_r(doc, "1. Timing sugerido (total ~15 min)", 1)
    add_bullets(
        doc,
        [
            "Contexto e problema (2 min)",
            "Objetivos e organização do trabalho (2 min)",
            "Pipeline e técnicas — Trilha A (4 min)",
            "Modelos e configuração (2 min)",
            "Avaliação: Recall@k e rubrica qualitativa (3 min)",
            "Resultados, observações e limitações (2 min)",
        ],
    )

    add_heading_r(doc, "2. Contexto", 1)
    add_p(
        doc,
        "O projeto implementa um sistema de consulta em linguagem natural sobre pareceres tributários "
        "da SEFAZ-GO (ICMS, FUNDEINFRA, regimes especiais), com documentos obtidos a partir do portal "
        "público de pareceres do Estado. O desafio é combinar recuperação de informação em texto "
        "jurídico em português com geração de respostas fundamentadas (RAG), evitando invenção de "
        "normas e permitindo verificação por citações aos trechos recuperados.",
    )

    add_heading_r(doc, "3. Objetivos", 1)
    add_bullets(
        doc,
        [
            "Construir pipeline reprodutível: coleta → transformação → ingestão semântica → indexação híbrida → chat.",
            "Cumprir a Trilha A: recuperação densa (embeddings), esparsa (BM25) e fusão por RRF.",
            "Medir qualidade do retriever com Recall@k (k ∈ {3, 5, 10}) e golden set com 29 perguntas válidas.",
            "Comparar dois encoders: MiniLM (geral) vs. modelo jurídico (BERT legal português).",
            "Avaliar qualitativamente respostas (groundedness, correção, citações, alucinação, recusa).",
            "Disponibilizar interface Streamlit para demonstrar os três modos de busca lado a lado.",
        ],
    )

    add_heading_r(doc, "4. Organização do desenvolvimento", 1)
    add_p(doc, "Fases técnicas (monorepositório Python 3.11):", bold=True)
    add_bullets(
        doc,
        [
            "Scraping: download controlado a partir da URL configurada; manifesto em `data/raw/`.",
            "Transform: conversão para Markdown (MarkItDown / LibreOffice para formatos legados); normalização opcional de títulos jurídicos.",
            "Ingest: chunking em dois níveis — secções (EMENTA, RELATÓRIO, FUNDAMENTAÇÃO, …) e janela de tokens (max_tokens ~700, overlap ~140); metadados e texto enriquecido para o denso.",
            "Index: embeddings no ChromaDB + índice BM25 persistido (`bm25_index.pkl`).",
            "Avaliação: `pipeline.py --step evaluate` e script `run_recall_at_k_values.py` para consolidar métricas em `docs/dados/recall_por_k.json`.",
            "Interface: `streamlit run app/app.py` com seleção Híbrido / Denso / Esparso.",
        ],
    )

    add_heading_r(doc, "5. Técnicas principais", 1)
    add_heading_r(doc, "5.1 Recuperação densa", 2)
    add_p(
        doc,
        "SentenceTransformers + ChromaDB: cada chunk (texto enriquecido com título, seção, normas, tributos) "
        "é vetorizado na indexação; na consulta, a pergunta é codificada no mesmo espaço e recuperam-se os "
        "vizinhos mais próximos.",
    )
    add_heading_r(doc, "5.2 Recuperação esparsa (BM25)", 2)
    add_p(
        doc,
        "BM25Okapi sobre tokens da query e do corpus de chunks; privilegia correspondência lexical (siglas, "
        "nomes de institutos, referências exatas). Não depende do modelo de embeddings — por isso os resultados "
        "\"esparsos\" são idênticos ao trocar apenas MiniLM por BERT legal.",
    )
    add_heading_r(doc, "5.3 Híbrido — Reciprocal Rank Fusion (RRF)", 2)
    add_p(
        doc,
        "Duas listas ranqueadas (denso e esparso) fundem-se somando 1/(k_rrf + rank + 1) por lista; o mesmo chunk_id "
        "aparece uma vez com score agregado (efeito de deduplicação lógica). Hiperparâmetro rrf_k = 60; top_k = 5 "
        "na avaliação típica; reranker opcional desativado neste trabalho.",
    )
    add_heading_r(doc, "5.4 Geração (LLM)", 2)
    add_p(
        doc,
        "LangChain com Ollama (ex.: Ministral) ou Google Gemini, conforme configuração. Prompt exige ancoragem nos "
        "trechos, citações no formato [[TRECHO_n]] e recusa padronizada quando não há suporte no contexto.",
    )

    add_heading_r(doc, "6. Modelos utilizados", 1)
    add_bullets(
        doc,
        [
            "Embeddings (comparação experimental): sentence-transformers/all-MiniLM-L6-v2 vs. stjiris/bert-large-portuguese-cased-legal-tsdae.",
            "LLM local: Ollama (ex.: ministral-3:8b); alternativa: API Google (Gemini) com chave em .env.",
            "Tokenização NLTK na divisão por tokens e na query BM25 (com fallback regex).",
        ],
    )

    add_heading_r(doc, "7. Avaliação", 1)
    add_p(doc, "Quantitativo:", bold=True)
    add_bullets(
        doc,
        [
            "Recall@k: documento esperado do golden presente em qualquer um dos top-k chunks (29 casos).",
            "Valores consolidados em docs/dados/recall_por_k.json (atualização de referência 2026-04-02).",
        ],
    )
    add_p(doc, "Qualitativo:", bold=True)
    add_bullets(
        doc,
        [
            "Rubrica (≥15 avaliações) com critérios: groundedness, correção, citações, alucinação, recusa.",
            "Exemplos: FundEinfra (BM25 forte), pergunta ANVISA fora do corpus (recusa no denso vs. alucinação no esparso).",
        ],
    )

    add_heading_r(doc, "8. Resultados numéricos (destaque — MiniLM, 29 casos)", 1)
    add_p(doc, "Recall por modo (valores do JSON consolidado):", bold=True)
    add_bullets(
        doc,
        [
            "k=3 — Denso 24,14% (7/29) | Esparso 65,52% (19/29) | Híbrido 55,17% (16/29).",
            "k=5 — Denso 27,59% (8/29) | Esparso 72,41% (21/29) | Híbrido 68,97% (20/29).",
            "k=10 — Denso 44,83% (13/29) | Esparso 75,86% (22/29) | Híbrido 82,76% (24/29).",
        ],
    )
    add_p(
        doc,
        "Interpretação: o esparso supera o denso MiniLM neste golden; o híbrido melhora o denso em todos os k; "
        "em k=3 e k=5 fica abaixo do só BM25; em k=10 supera denso e esparso e o denso sobe face a k=5.",
    )
    add_p(doc, "BERT legal no denso:", bold=True)
    add_bullets(
        doc,
        [
            "Denso piora frente ao MiniLM em k=3 e k=5; esparso idêntico entre modelos (BM25).",
            "Híbrido: Recall@5 BERT 72,41% (21/29) vs. MiniLM 68,97% (20/29); em k=10 ambos 82,76% (24/29).",
        ],
    )

    add_heading_r(doc, "9. Observações pertinentes", 1)
    add_bullets(
        doc,
        [
            "Concorrência GPU: embeddings (PyTorch) e Ollama podem disputar VRAM — documentado no README (CPU nos embeddings na demo quando necessário).",
            "Latência: prompts longos (vários chunks) + duas passagens LLM possíveis (correção de citações) aumentam tempo de resposta.",
            "Chunking por token pode cortar listas ou artigos ao meio; trade-off por tamanho estável de contexto.",
            "Fora do corpus: métrica automática de Recall ignora linhas sem gabarito; rubrica e interface testam recusa e alucinação.",
        ],
    )

    add_heading_r(doc, "10. Encerramento e demonstração", 1)
    add_bullets(
        doc,
        [
            "Mensagem central: RAG híbrido é adequado a pareceres com léxico técnico; RRF recupera falhas do denso isolado neste conjunto.",
            "Demo sugerida: mesma pergunta nos três modos + painel de trechos recuperados.",
            "Reprodutibilidade: config.yaml + pipeline documentado no README e RELATORIO_PR7.md.",
        ],
    )

    add_heading_r(doc, "Referências rápidas no repositório", 1)
    add_bullets(
        doc,
        [
            "RELATORIO_PR7.md — relatório técnico completo.",
            "README.md — instalação, etapas do pipeline, GPU/CPU.",
            "app/app.py — interface; src/rag_core.py — retrieve + LLM; src/pipeline.py — ingest/index.",
        ],
    )

    doc.save(OUT)
    print(f"Escrito: {OUT}")


if __name__ == "__main__":
    main()
